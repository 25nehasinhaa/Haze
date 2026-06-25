"""
Converts real challenge JSONL candidates → internal CandidateProfile dicts.
Also parses job_description.docx → JobDescription dict.
All conversion happens here; downstream scoring is unchanged.
"""

from __future__ import annotations

import json
from datetime import datetime, date
from pathlib import Path
from typing import Iterator

from backend.app.schemas import (
    CandidateProfile, CareerEvent, EvidenceItem, JobDescription
)

# ---------------------------------------------------------------------------
# Skill / evidence extraction from JSONL candidate
# ---------------------------------------------------------------------------

PROFICIENCY_WEIGHTS = {
    "expert": "strong",
    "advanced": "strong",
    "intermediate": "partial",
    "beginner": "weak",
}

# Redrob assessment score → evidence level
def _assessment_level(score: float) -> str:
    if score >= 75:
        return "strong"
    if score >= 50:
        return "partial"
    if score >= 25:
        return "weak"
    return "unsupported"


def _months_ago_year(months: int) -> int:
    """Convert duration_months-ago into an approximate year."""
    current_year = datetime.now().year
    return max(current_year - months // 12, 1990)


def _recency_year_from_career(career_history: list[dict]) -> int:
    """Latest end_date year across career history (or current year if current role)."""
    latest = datetime.now().year
    for role in career_history:
        if role.get("is_current"):
            return datetime.now().year
        end = role.get("end_date")
        if end:
            try:
                return int(str(end)[:4])
            except Exception:
                pass
    return latest


def jsonl_to_internal(record: dict) -> dict:
    """Convert one JSONL candidate record to internal CandidateProfile dict."""
    profile = record.get("profile", {})
    career = record.get("career_history", [])
    skills_raw = record.get("skills", [])
    certs = record.get("certifications", [])
    redrob = record.get("redrob_signals", {})
    assessment_scores: dict[str, float] = redrob.get("skill_assessment_scores", {})

    name = profile.get("anonymized_name", record.get("candidate_id", "Unknown"))
    headline = profile.get("headline", "")
    summary = profile.get("summary", "")
    yoe = float(profile.get("years_of_experience", 0))
    current_year = datetime.now().year

    # --- Skills list (names only) ---
    skill_names = [s["name"] for s in skills_raw if s.get("name")]

    # --- Evidence items ---
    evidence: list[dict] = []

    # From skills: use proficiency + assessment score if available
    for skill in skills_raw:
        sname = skill.get("name", "")
        if not sname:
            continue
        proficiency = skill.get("proficiency", "intermediate")
        duration_months = skill.get("duration_months", 12)
        endorsements = skill.get("endorsements", 0)
        last_used_year = _months_ago_year(max(0, int(yoe * 12) - duration_months))

        # Override with Redrob assessment if present
        if sname in assessment_scores:
            level = _assessment_level(assessment_scores[sname])
        else:
            level = PROFICIENCY_WEIGHTS.get(proficiency, "partial")
            # Boost with endorsements
            if endorsements >= 20 and level == "partial":
                level = "strong"
            elif endorsements == 0 and level in ("strong", "partial"):
                level = "weak"

        source_parts = []
        if endorsements > 0:
            source_parts.append(f"{endorsements} endorsements")
        if sname in assessment_scores:
            source_parts.append(f"assessment score {assessment_scores[sname]:.0f}/100")
        if duration_months:
            source_parts.append(f"{duration_months} months experience")
        source = f"{sname}: " + "; ".join(source_parts) if source_parts else sname

        evidence.append({
            "claim": sname,
            "level": level,
            "last_used": last_used_year,
            "source": source,
        })

    # From career history descriptions: extract behavioral evidence
    behavioral_markers = {
        "led": "led", "owned": "owned", "deployed": "deployed",
        "built": "built", "shipped": "shipped", "improved": "improved",
        "mentored": "mentored", "collaborated": "collaborated",
        "stakeholder": "stakeholder", "production": "production deployment",
        "launched": "launched", "reduced": "reduced",
    }
    seen_behaviors: set[str] = set()
    for role in career:
        desc = role.get("description", "").lower()
        role_year = current_year
        if role.get("end_date") and not role.get("is_current"):
            try:
                role_year = int(str(role["end_date"])[:4])
            except Exception:
                pass
        for marker, label in behavioral_markers.items():
            if marker in desc and label not in seen_behaviors:
                seen_behaviors.add(label)
                # Find snippet
                idx = desc.find(marker)
                snippet = role.get("description", "")[max(0, idx-60):idx+120].strip()
                evidence.append({
                    "claim": label,
                    "level": "partial",
                    "last_used": role_year,
                    "source": snippet[:200],
                })

    # From certifications
    for cert in certs:
        cname = cert.get("name", "")
        year = cert.get("year", current_year - 2)
        if cname:
            evidence.append({
                "claim": cname,
                "level": "strong",
                "last_used": year,
                "source": f"Certified by {cert.get('issuer', 'issuer')} in {year}",
            })

    # GitHub as signal
    github_score = redrob.get("github_activity_score", -1)
    if github_score and github_score > 0:
        level = "strong" if github_score >= 60 else "partial" if github_score >= 30 else "weak"
        evidence.append({
            "claim": "open source contributions",
            "level": level,
            "last_used": current_year,
            "source": f"GitHub activity score: {github_score}/100",
        })

    # --- Domains from career industries ---
    INDUSTRY_TO_DOMAIN = {
        "fintech": "fintech", "financial": "fintech", "banking": "banking",
        "payments": "payments", "insurance": "risk", "risk": "risk",
        "healthcare": "healthcare", "health": "healthcare",
        "ecommerce": "ecommerce", "retail": "retail", "e-commerce": "ecommerce",
        "iot": "iot", "cloud": "cloud", "security": "security",
        "recruiting": "recruiting", "hr": "hr", "human resources": "hr",
        "saas": "cloud", "ai": "ai", "ml": "ai", "data": "data",
        "it services": "cloud", "software": "cloud",
    }
    domains: set[str] = set()
    for role in career:
        industry = role.get("industry", "").lower()
        for keyword, domain in INDUSTRY_TO_DOMAIN.items():
            if keyword in industry:
                domains.add(domain)
    current_industry = profile.get("current_industry", "").lower()
    for keyword, domain in INDUSTRY_TO_DOMAIN.items():
        if keyword in current_industry:
            domains.add(domain)

    # --- Career events from history ---
    career_events: list[dict] = []
    for role in sorted(career, key=lambda r: r.get("start_date", ""), reverse=False):
        try:
            year = int(str(role.get("start_date", ""))[:4])
        except Exception:
            year = current_year - 2
        title = role.get("title", "")
        company = role.get("company", "")
        career_events.append({
            "year": year,
            "event": f"{title} at {company}",
        })

    # Fallback if no history
    if not career_events:
        career_events.append({
            "year": current_year - max(1, int(yoe)),
            "event": profile.get("current_title", "Professional"),
        })

    # --- Redrob behavioral signals as extra evidence ---
    # Availability / engagement signals baked into dedicated redrob_signals field
    # We store them separately so signalrank can access them
    redrob_evidence = _build_redrob_evidence(redrob, current_year)

    return {
        "candidate_id": record.get("candidate_id", ""),
        "name": name,
        "headline": headline,
        "summary": summary,
        "skills": skill_names,
        "domains": list(domains),
        "evidence": evidence + redrob_evidence,
        "career_events": career_events,
        # Preserve raw redrob signals for availability scoring
        "redrob_signals": redrob,
        "years_of_experience": yoe,
        "current_title": profile.get("current_title", ""),
        "location": profile.get("location", ""),
        "country": profile.get("country", ""),
    }


def _build_redrob_evidence(redrob: dict, current_year: int) -> list[dict]:
    """Convert redrob platform signals to evidence items."""
    evidence = []

    # Verified identity
    if redrob.get("verified_email") and redrob.get("verified_phone"):
        evidence.append({
            "claim": "verified profile",
            "level": "strong",
            "last_used": current_year,
            "source": "Email and phone verified on Redrob platform",
        })

    # High response rate
    rr = redrob.get("recruiter_response_rate", 0)
    if rr >= 0.7:
        evidence.append({
            "claim": "high recruiter responsiveness",
            "level": "strong",
            "last_used": current_year,
            "source": f"Recruiter response rate: {rr:.0%}",
        })
    elif rr <= 0.1:
        evidence.append({
            "claim": "low recruiter responsiveness",
            "level": "unsupported",
            "last_used": current_year,
            "source": f"Recruiter response rate: {rr:.0%} — candidate may be unresponsive",
        })

    # Active recently
    last_active = redrob.get("last_active_date", "")
    if last_active:
        try:
            la_date = date.fromisoformat(str(last_active))
            days_inactive = (date.today() - la_date).days
            if days_inactive <= 30:
                evidence.append({
                    "claim": "recently active",
                    "level": "strong",
                    "last_used": current_year,
                    "source": f"Last active {days_inactive} days ago",
                })
            elif days_inactive > 180:
                evidence.append({
                    "claim": "inactive profile",
                    "level": "unsupported",
                    "last_used": current_year - 1,
                    "source": f"Last active {days_inactive} days ago — low availability signal",
                })
        except Exception:
            pass

    return evidence


# ---------------------------------------------------------------------------
# JD parsing from DOCX
# ---------------------------------------------------------------------------

def parse_jd_docx(docx_path: Path) -> dict:
    """Extract job description from DOCX into internal JD dict format."""
    from docx import Document

    doc = Document(docx_path)
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    # Extract title from first non-empty line
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    title = lines[0] if lines else "Hiring Role"
    # Clean "Job Description: " prefix
    if ":" in title and len(title.split(":")[0]) < 30:
        title = title.split(":", 1)[1].strip()

    return _extract_jd_from_text(title, full_text)


def _extract_jd_from_text(title: str, text: str) -> dict:
    """Parse structured JD fields from raw text."""
    import re

    text_lower = text.lower()

    # Core ML/AI skills — use names that match the actual dataset skill names
    MUST_HAVE_CANDIDATES = [
        ("embeddings", ["embedding", "embeddings"]),
        ("semantic search", ["semantic search"]),
        ("vector search", ["vector search", "vector database", "vector db", "faiss", "pinecone", "qdrant", "weaviate", "milvus", "pgvector", "opensearch", "elasticsearch"]),
        ("information retrieval", ["retrieval", "information retrieval", "bm25", "hybrid search"]),
        ("python", ["python"]),
        ("recommendation systems", ["recommendation", "ranking", "rerank"]),
        ("sentence transformers", ["sentence-transformers", "sentence transformers", "bge", "e5 model"]),
        ("fine-tuning llms", ["fine-tun", "lora", "qlora", "peft", "finetuning"]),
        ("nlp", ["nlp", "natural language processing"]),
        ("learning to rank", ["learning to rank", "learning-to-rank", "ndcg", "mrr", "evaluation framework", "offline eval"]),
        ("langchain", ["langchain", "llm framework", "rag"]),
        ("faiss", ["faiss", "ann", "approximate nearest"]),
    ]
    NICE_TO_HAVE_CANDIDATES = [
        ("hugging face transformers", ["hugging face", "transformers library"]),
        ("elasticsearch", ["elasticsearch", "opensearch"]),
        ("machine learning", ["machine learning", "ml engineering"]),
        ("a/b testing", ["a/b test", "ab test", "online eval"]),
        ("xgboost", ["xgboost", "gradient boost", "lgbm"]),
        ("llms", ["llm", "large language model", "gpt", "openai"]),
    ]

    must_have = []
    for canonical, patterns in MUST_HAVE_CANDIDATES:
        if any(p in text_lower for p in patterns):
            must_have.append(canonical)
    must_have = must_have[:10]

    nice_to_have = []
    for canonical, patterns in NICE_TO_HAVE_CANDIDATES:
        if any(p in text_lower for p in patterns):
            nice_to_have.append(canonical)
    nice_to_have = nice_to_have[:6]

    if not must_have:
        must_have = ["python", "machine learning", "nlp", "information retrieval", "recommendation systems"]

    # Domains
    DOMAIN_CANDIDATES = [
        "recruiting", "hr-tech", "talent", "fintech", "payments", "ai", "ml",
        "search", "recommendation", "startup",
    ]
    domains = [d for d in DOMAIN_CANDIDATES if d in text_lower]

    # Seniority signals
    seniority_signals = []
    for signal in ["production deployment", "end-to-end", "mentoring", "architecture",
                   "cross-functional", "evaluation", "a/b test", "own the"]:
        if signal in text_lower:
            seniority_signals.append(signal)

    return {
        "title": title,
        "summary": text[:600].replace("\n", " "),
        "must_have_skills": must_have,
        "nice_to_have_skills": nice_to_have,
        "domains": domains[:6],
        "seniority_signals": seniority_signals[:8],
    }


# ---------------------------------------------------------------------------
# JSONL streaming loader
# ---------------------------------------------------------------------------

def stream_jsonl(path: Path) -> Iterator[dict]:
    """Stream raw JSONL records one by one."""
    with path.open("r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if line:
                try:
                    yield json.loads(line)
                except json.JSONDecodeError:
                    continue


def load_jsonl_candidates(
    path: Path,
    limit: int | None = None,
    filter_open_to_work: bool = False,
) -> list[dict]:
    """Load and convert JSONL candidates to internal schema dicts."""
    results = []
    for i, record in enumerate(stream_jsonl(path)):
        if limit and i >= limit:
            break
        if filter_open_to_work:
            signals = record.get("redrob_signals", {})
            if not signals.get("open_to_work_flag", False):
                continue
        results.append(jsonl_to_internal(record))
    return results


# ---------------------------------------------------------------------------
# Process-level cache: avoid re-loading + re-converting 100k records and
# rebuilding the TF-IDF index on every request. Cache is keyed by
# (path, limit, filter_open_to_work) and invalidated only if those change.
# ---------------------------------------------------------------------------

_CANDIDATE_CACHE: dict[tuple, list[dict]] = {}


def load_jsonl_candidates_cached(
    path: Path,
    limit: int | None = None,
    filter_open_to_work: bool = False,
) -> list[dict]:
    key = (str(path), limit, filter_open_to_work)
    if key not in _CANDIDATE_CACHE:
        _CANDIDATE_CACHE[key] = load_jsonl_candidates(path, limit, filter_open_to_work)
    return _CANDIDATE_CACHE[key]


def clear_candidate_cache() -> None:
    _CANDIDATE_CACHE.clear()


# ---------------------------------------------------------------------------
# CSV submission exporter
# ---------------------------------------------------------------------------

def export_submission_csv(ranking: dict, output_path: Path, top_n: int = 100) -> Path:
    """
    Export ranked candidates to submission CSV format:
    candidate_id,rank,score,reasoning
    - Exactly 100 rows
    - Scores non-increasing by rank
    - Tie-break: candidate_id ascending
    - No duplicate candidate_ids
    """
    import csv

    candidates = ranking["candidates"][:top_n]
    n = len(candidates)

    if n == 0:
        raise ValueError("No candidates to export")

    # Normalize scores into 0.200..0.992 range
    raw_scores = [c["signal_score"] for c in candidates]
    max_s = max(raw_scores)
    min_s = min(raw_scores)
    score_range = max(max_s - min_s, 0.001)

    def normalize(s: float) -> float:
        normalized = (s - min_s) / score_range
        return round(0.200 + normalized * 0.792, 4)

    # Assign normalized scores, round to 4dp, handle ties via candidate_id sort
    rows = []
    for c in candidates:
        cid = c.get("candidate_id") or c["name"]
        score = normalize(c["signal_score"])
        trust = c.get("trust_label", "")
        top_strength = c.get("strengths", ["Matched candidate"])[0]
        yoe = c.get("years_of_experience", "")
        title = c.get("current_title", c["name"])
        reasoning = f"{title} with {yoe} yrs; {trust}; {top_strength[:60]}" if yoe else f"{trust}; {top_strength[:80]}"
        rows.append({"candidate_id": cid, "score": score, "reasoning": reasoning[:200]})

    # Sort: primary = score descending, secondary = candidate_id ascending (tie-break per spec)
    rows.sort(key=lambda r: (-r["score"], r["candidate_id"]))

    # Pad or trim to exactly top_n rows (shouldn't be needed but safety)
    rows = rows[:top_n]

    output_path.parent.mkdir(parents=True, exist_ok=True)
    with output_path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["candidate_id", "rank", "score", "reasoning"])
        for i, row in enumerate(rows, start=1):
            writer.writerow([row["candidate_id"], i, f"{row['score']:.4f}", row["reasoning"]])

    return output_path


# ---------------------------------------------------------------------------
# Availability scorer (redrob signals → multiplier)
# ---------------------------------------------------------------------------

def compute_availability_score(redrob: dict) -> float:
    """
    Compute 0..1 availability multiplier from redrob signals.
    Applied as a post-scoring modifier, not mixed into signal_score directly.
    """
    score = 0.5  # baseline

    # Open to work
    if redrob.get("open_to_work_flag"):
        score += 0.15

    # Response rate
    rr = redrob.get("recruiter_response_rate", 0.3)
    score += 0.15 * rr

    # Recently active (last 90 days = good)
    last_active = redrob.get("last_active_date", "")
    if last_active:
        try:
            la_date = date.fromisoformat(str(last_active))
            days = (date.today() - la_date).days
            if days <= 30:
                score += 0.15
            elif days <= 90:
                score += 0.08
            elif days > 180:
                score -= 0.15
        except Exception:
            pass

    # Notice period
    notice = redrob.get("notice_period_days", 60)
    if notice <= 15:
        score += 0.10
    elif notice <= 30:
        score += 0.05
    elif notice > 90:
        score -= 0.05

    # Interview completion
    icr = redrob.get("interview_completion_rate", 0.5)
    score += 0.05 * icr

    return max(0.1, min(1.0, score))
